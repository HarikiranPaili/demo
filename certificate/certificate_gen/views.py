import datetime

from django.shortcuts import render,redirect
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from io import BytesIO
from openpyxl.comments import Comment
import os
import io
from openpyxl.styles import PatternFill
from openpyxl.comments import Comment
from openpyxl.utils.cell import get_column_letter
from openpyxl.styles import Color, PatternFill, Font, Alignment
from openpyxl.utils import quote_sheetname
from openpyxl.worksheet.cell_range import CellRange
from openpyxl.utils import absolute_coordinate
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Font, Color, colors, Alignment, NamedStyle, Border, Side, PatternFill, Protection, Color
from .models import *
from django.core.files.base import ContentFile
import docx
import openpyxl
from docx.shared import RGBColor
import io
from django.core.files import File
from django.conf import settings
from django.core.files.storage import default_storage


# relative import of forms

# importing formset_factory

def text(request):
    context = {}
    if request.method == 'POST':
        wordsfile = request.FILES['wordsfile']
        text = request.POST.get('text')
        context['text'] = text
        workbook = openpyxl.load_workbook(wordsfile)
        worksheet = workbook.active
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            # Get the values from the current row
            cell1_value = row[0]
            cell2_value = row[1]
            if cell1_value in text:
                text = text.replace(cell1_value, cell2_value)
                context['convertedtext'] = text
    return render(request, "text.html",context)


def file(request):
    context = {}
    data = Files.objects.all()
    context['data'] = data
    if request.method == 'POST':
        file = request.FILES['file']
        wordsfile = request.FILES['wordsfile']
        if str(file).endswith(".docx"):
            # Open the Word file
            doc = docx.Document(file)
            # Open the Keywords file
            wb = openpyxl.load_workbook(wordsfile)
            ws = wb.active

            # Loop through each row of the Keywords file
            for row in ws.iter_rows(min_row=2, values_only=True):
                keyword = row[0]
                alt_word = row[1]

                # Loop through each paragraph in the Word file
                for para in doc.paragraphs:
                    # Check if the paragraph contains the keyword
                    if keyword in para.text:
                        # Add the error word and suggestion next to the keyword
                        para.add_run(f" ^Error: consider using '{alt_word}'")

                        # Highlight the error word in red
                        for run in para.runs:
                            if '^Error' in run.text:
                                font = run.font
                                font.color.rgb = RGBColor(255, 0, 0)

            # Save the reviewed Word file on the desktop
            buffer = io.BytesIO()
            doc.save(buffer)
            content = ContentFile(buffer.getvalue())
            if content:
                processed_file = Files.objects.create(
                    time=str(datetime.datetime.now().replace(second=0, microsecond=0)))
                processed_file.documents.save('processed_' + str(file), content)

            #doc.save('C:/Users/rpaili/Desktop/Pro_Dummy.docx')
        else:
            workbook = openpyxl.load_workbook(file)

            # Load the Keywords sheet
            workbook_1 = openpyxl.load_workbook(wordsfile)
            keywords_sheet = workbook_1['Sheet1']

            # Create a dictionary to map each keyword to its alternatives
            keywords_dict = {}
            for row in keywords_sheet.iter_rows(min_row=2, values_only=True):
                keyword = row[0]
                alternatives = row[1].split(', ') if row[1] else []
                keywords_dict[keyword] = alternatives

            # Select the first worksheet
            worksheet = workbook.active

            # Loop through each cell in the worksheet
            for row in worksheet.iter_rows():
                for cell in row:
                    # Check for keywords and highlight the cell if found
                    if cell.value:
                        value = str(cell.value) if not isinstance(cell.value, str) else cell.value
                        words = value.split()
                        matched_keywords = set()
                        for i, word in enumerate(words):
                            for keyword, alternatives in keywords_dict.items():
                                if keyword == word:
                                    # if keyword in word:
                                    # Highlight the keyword with yellow background color and red font color
                                    cell.fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')

                                    # Add the "(^Error)" word after the identified keyword
                                    keyword_position = word.find(keyword)
                                    keyword_length = len(keyword)
                                    error_message = "(^Error)"
                                    if keyword_position > 0 and word[keyword_position - 1].isalpha():
                                        # if word[keyword_position-1].isalpha():
                                        # If the keyword is preceded by an alphabet, add the error message with a space
                                        error_message = " " + error_message
                                    words[i] = word[:keyword_position] + keyword + error_message + word[
                                                                                                   keyword_position + keyword_length:]

                                    # Add the matched keyword to the set of matched keywords
                                    matched_keywords.add(keyword)

                                    # Construct the comment text to include alternatives for all matched keywords
                                    comment_text = ''
                                    i = 1
                                    for keyword in matched_keywords:
                                        alternatives = keywords_dict[keyword]
                                        if alternatives:
                                            comment_text += f"{i}. '{keyword}': {', '.join(alternatives)}\n"
                                            i += 1
                                    if comment_text:
                                        comment_text = "Consider alternatives for:\n" + comment_text.strip()

                                        # Add comment to the cell
                                        comment = Comment(comment_text, "Comment Author")
                                        cell.comment = comment

                                    value = " ".join(words)
                                    cell.value = value

            # Save the processed workbook to a Django FileField
            buffer = io.BytesIO()
            workbook.save(buffer)
            content = ContentFile(buffer.getvalue())
            if content:
                processed_file = Files.objects.create(
                    time=str(datetime.datetime.now().replace(second=0, microsecond=0)))
                processed_file.documents.save('processed_' + str(file), content)

        #file_stream = io.BytesIO(file)
        return redirect('file')
    return render(request,'file.html',context)

# import ipywidgets as widgets
#         import openpyxl
#         from openpyxl.styles import Font, PatternFill, Alignment
#         from io import BytesIO
#         from openpyxl.comments import Comment
#         import os
#         import io
#         from openpyxl.styles import PatternFill
#         from openpyxl.comments import Comment
#         from openpyxl.utils.cell import get_column_letter
#         from openpyxl.styles import Color, PatternFill, Font, Alignment
#         from openpyxl.utils import quote_sheetname
#         from openpyxl.worksheet.cell_range import CellRange
#         from openpyxl.utils import absolute_coordinate
#         from openpyxl.utils import get_column_letter
#         from openpyxl.styles import PatternFill, Font, Color, colors, Alignment, NamedStyle, Border, Side, PatternFill, Protection, Color
#
#         # Create GUI buttons
#         from IPython.display import display
#
#         # Create the output widget
#         processed_output = widgets.Output()
#
#         # Create the GUI widgets
#         upload_button = widgets.FileUpload(accept='.xlsx', multiple=False)
#         process_button = widgets.Button(description='Process')
#         save_button = widgets.Button(description='Save')
#         clear_button = widgets.Button(description='Clear')
#         message_label = widgets.Label()
#
#         # Display the widgets
#         display(widgets.VBox([widgets.HBox([upload_button, process_button, save_button, clear_button]), message_label, processed_output]))
#
#         # Function to process the Excel file
#         def process_excel(file_contents):
#             # Load the workbook from the file contents
#             file_stream = io.BytesIO(file_contents)
#             workbook = openpyxl.load_workbook(file_stream)
#
#             # Load the Keywords sheet
#             file_path = 'C:/Users/rpaili/Desktop/Keywords.xlsx'
#             workbook_1 = openpyxl.load_workbook(file_path)
#             keywords_sheet = workbook_1['Sheet1']
#
#             # Create a dictionary to map each keyword to its alternatives
#             keywords_dict = {}
#             for row in keywords_sheet.iter_rows(min_row=2, values_only=True):
#                 keyword = row[0]
#                 alternatives = row[1].split(', ') if row[1] else []
#                 keywords_dict[keyword] = alternatives
#
#             # Select the first worksheet
#             worksheet = workbook.active
#
#             # Loop through each cell in the worksheet
#             for row in worksheet.iter_rows():
#                 for cell in row:
#                     # Check for keywords and highlight the cell if found
#                     if cell.value:
#                         value = str(cell.value) if not isinstance(cell.value, str) else cell.value
#                         words = value.split()
#                         matched_keywords = set()
#                         for i, word in enumerate(words):
#                             for keyword, alternatives in keywords_dict.items():
#                                 if keyword in word:
#                                     # Highlight the keyword with yellow background color and red font color
#                                     cell.fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
#
#                                     # Add the "(^Error)" word after the identified keyword
#                                     keyword_position = word.find(keyword)
#                                     keyword_length = len(keyword)
#                                     error_message = "(^Error)"
#                                     if keyword_position > 0 and word[keyword_position-1].isalpha():
#                                         # If the keyword is preceded by an alphabet, add the error message with a space
#                                         error_message = " " + error_message
#                                     words[i] = word[:keyword_position] + keyword + error_message + word[keyword_position+keyword_length:]
#
#                                     # Add the matched keyword to the set of matched keywords
#                                     matched_keywords.add(keyword)
#
#                                     # Construct the comment text to include alternatives for all matched keywords
#                                     comment_text = ''
#                                     i = 1
#                                     for keyword in matched_keywords:
#                                         alternatives = keywords_dict[keyword]
#                                         if alternatives:
#                                             comment_text += f"{i}. '{keyword}': {', '.join(alternatives)}\n"
#                                             i += 1
#                                     if comment_text:
#                                         comment_text = "Consider alternatives for:\n" + comment_text.strip()
#
#                                         # Add comment to the cell
#                                         comment = Comment(comment_text, "Comment Author")
#                                         cell.comment = comment
#
#                                     value = " ".join(words)
#                                     cell.value = value
#
#             # Return the modified workbook
#             return workbook
#
#         # Function to save the processed Excel file to the desktop
#         def save_excel(_):
#             # Get the path to the desktop directory
#             desktop_path = os.path.join(os.environ['USERPROFILE'], 'Desktop')
#             # Get the filename from the uploaded file
#             filename = next(iter(upload_button.value))
#             # Process the uploaded workbook
#             processed_workbook = process_excel(upload_button.value[next(iter(upload_button.value))]['content'])
#             # Save the processed workbook to the desktop directory with a '_processed' suffix
#             processed_workbook.save(os.path.join(desktop_path, f'{filename}_processed.xlsx'))
#             # Display a message after saving is complete
#             message_label.value = 'File saved to desktop.'
#
#         # Function to clear the GUI widgets and output
#         def clear_widgets(_):
#             upload_button.value.clear()
#             upload_button._counter = 0
#             processed_output.clear_output()
#             message_label.value = ''
#
#         # Function to handle process button click event
#         def on_process_button_clicked(_):
#             with processed_output:
#                 processed_output.clear_output
#                 # Process the uploaded workbook
#                 process_excel(upload_button.value[next(iter(upload_button.value))]['content'])
#                 # Display a message after processing is complete
#                 print('Dictionary Checked.')
#                 print('Vocabulary Checked.')
#                 print('Grammer issues Checked.')
#                 print('Case Sensitivity Checked.')
#                 print('Processing Completed.')
#
#         # Attach the functions to the GUI buttons
#         process_button.on_click(on_process_button_clicked)
#         save_button.on_click(save_excel)
#         clear_button.on_click(clear_widgets)


